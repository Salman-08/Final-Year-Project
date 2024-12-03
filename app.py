from flask import Flask, request, jsonify
from flask_cors import CORS
import docx
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer, util
from mysql.connector import pooling, Error
import logging
import re
import os
from docx.enum.text import WD_COLOR_INDEX
from concurrent.futures import ThreadPoolExecutor, as_completed

# Configure logging
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Directory where reports will be saved
REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# Database connection pooling
dbconfig = {
    "host": "localhost",
    "user": "root",
    "password": "",
    "database": "aae"
}
connection_pool = pooling.MySQLConnectionPool(pool_name="mypool",
                                              pool_size=10,
                                              **dbconfig)

# Load models outside of route functions to avoid reloading them on each request
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""
    with open(file_path, 'rb') as f:
        reader = PdfReader(f)
        text = "\n".join([page.extract_text() for page in reader.pages])
    return text.strip()

def normalize_text(text):
    """Normalize text by converting to lowercase and removing extra spaces."""
    return re.sub(r'\s+', ' ', text.lower()).strip()

def get_submission_text(file_path):
    """Get text from submission files based on their format."""
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file format. Only DOCX and PDF files are supported.")

def evaluate_submission(teacher_text, student_text):
    """Evaluate submissions against the teacher's solution using semantic similarity."""
    teacher_text_normalized = normalize_text(teacher_text)
    student_text_normalized = normalize_text(student_text)

    # Encode the texts using Sentence-BERT
    teacher_embeddings = sbert_model.encode([teacher_text_normalized], convert_to_tensor=True)
    student_embeddings = sbert_model.encode([student_text_normalized], convert_to_tensor=True)

    # Compute the cosine similarity between the teacher and student embeddings
    similarity_scores = util.pytorch_cos_sim(teacher_embeddings, student_embeddings).item()

    return similarity_scores

def save_grade_to_db(submission_id, user_id, student_name, grade, similarity_score, report_path):
    """Save the grades and similarity scores to the database."""
    try:
        conn = connection_pool.get_connection()
        cursor = conn.cursor()
        query = "INSERT INTO grades (submission_id, user_id, student_name, grade, similarity_score, report_path) VALUES (%s, %s, %s, %s, %s, %s)"
        cursor.execute(query, (submission_id, user_id, student_name, grade, similarity_score, report_path))
        conn.commit()
    except Error as err:
        logging.error(f"Database error: {err}")
        raise
    finally:
        cursor.close()
        conn.close()

def highlight_differences_docx(teacher_text, student_text, output_path, grade, similarity_score):
    """Highlight differences in a DOCX file with comments explaining each change."""
    teacher_words = teacher_text.split()
    student_words = student_text.split()

    doc = docx.Document()

    doc.add_heading('Evaluation Report', level=1)
    doc.add_paragraph(f'Grade: {grade}')
    doc.add_paragraph(f'Similarity Score: {similarity_score:.2f}')
    doc.add_paragraph('\n')

    teacher_para = doc.add_paragraph()
    student_para = doc.add_paragraph()

    for word in teacher_words:
        run = teacher_para.add_run(word + ' ')
        if word not in student_words:
            run.font.highlight_color = WD_COLOR_INDEX.RED

    for word in student_words:
        run = student_para.add_run(word + ' ')
        if word not in teacher_words:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_paragraph('\n')
    doc.add_heading('Explanations for Highlights', level=2)
    doc.add_paragraph('Words missing in the student\'s submission are highlighted in red.')
    doc.add_paragraph('Words added by the student are highlighted in yellow.')

    doc.add_paragraph('\n')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(f'The student has received a grade of {grade} with a similarity score of {similarity_score:.2f}. Major strengths and weaknesses identified in the comparison.')

    doc.save(output_path)

def check_if_evaluated(submission_id):
    """Check if the submission has already been evaluated."""
    try:
        conn = connection_pool.get_connection()
        cursor = conn.cursor(dictionary=True)
        query = "SELECT * FROM grades WHERE submission_id = %s"
        cursor.execute(query, (submission_id,))
        result = cursor.fetchone()
        return result is not None
    except Error as err:
        logging.error(f"Database error: {err}")
        raise
    finally:
        cursor.close()
        conn.close()

def process_submission(submission_id):
    try:
        logging.debug(f"Processing submission ID {submission_id}")
        if check_if_evaluated(submission_id):
            logging.debug(f"Submission ID {submission_id} has already been evaluated.")
            return {"message": f"Submission ID {submission_id} has already been evaluated."}

        conn = connection_pool.get_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT * FROM submissions WHERE submission_id = %s", (submission_id,))
        submission = cursor.fetchone()
        
        if not submission:
            logging.error(f"Submission not found for ID {submission_id}.")
            return {"error": f"Submission not found for ID {submission_id}."}
        
        student_text = get_submission_text(submission['file_path'])
        
        assignment_id = submission['assignment_id']
        cursor.execute("SELECT * FROM solutions WHERE assignment_id = %s", (assignment_id,))
        solution = cursor.fetchone()
        
        if not solution:
            logging.error(f"Solution not found for assignment ID {assignment_id}.")
            return {"error": f"Solution not found for assignment ID {assignment_id}."}

        cursor.execute("SELECT initial_grade FROM assignments WHERE assignment_id = %s", (assignment_id,))
        assignment = cursor.fetchone()
        
        if not assignment:
            logging.error(f"Assignment not found for ID {assignment_id}.")
            return {"error": f"Assignment not found for ID {assignment_id}."}

        initial_grade = assignment['initial_grade']

        teacher_text = get_submission_text(solution['file_path'])
        
        similarity_score = evaluate_submission(teacher_text, student_text)
        
        grade_mapping = [
            (0.9, 0.9 * initial_grade),
            (0.8, 0.8 * initial_grade),
            (0.7, 0.7 * initial_grade),
            (0.6, 0.6 * initial_grade),
            (0.5, 0.5 * initial_grade),
            (0.4, 0.4 * initial_grade),
            (0.3, 0.3 * initial_grade),
            (0.2, 0.2 * initial_grade),
            (0.1, 0.1 * initial_grade),
            (0, 0)
        ]

        grade = next((g for score, g in grade_mapping if similarity_score >= score), 0)
        grade = round(grade)

        user_id = submission['user_id']
        
        cursor.execute("SELECT fullname FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            logging.error(f"User not found for ID {user_id}.")
            return {"error": f"User not found for ID {user_id}."}
        
        student_name = user['fullname']
        
        report_path = os.path.join(REPORTS_DIR, f"highlighted_report_{submission_id}.docx")
        highlight_differences_docx(teacher_text, student_text, report_path, grade, similarity_score)
        
        save_grade_to_db(submission_id, user_id, student_name, grade, similarity_score, report_path)
        
        return {'grade': grade, 'similarity_score': similarity_score, 'report_path': report_path}
    
    except Exception as e:
        logging.error(f"Error processing submission ID {submission_id}: {str(e)}")
        return {"error": str(e)}
    
    finally:
        cursor.close()
        conn.close()

@app.route('/evaluate', methods=['POST', 'OPTIONS'])
def evaluate():
    logging.debug("Received a POST request to /evaluate")
    
    if request.method == 'OPTIONS':
        # CORS preflight request
        return jsonify({'status': 'OK'}), 200

    data = request.json
    submission_ids = data.get('submission_ids')

    if not submission_ids or not isinstance(submission_ids, list):
        logging.error("Submission IDs are required and should be a list.")
        return jsonify({"error": "Submission IDs are required and should be a list."}), 400

    results = {}
    already_evaluated_messages = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(process_submission, submission_id): submission_id for submission_id in submission_ids}
        for future in as_completed(futures):
            submission_id = futures[future]
            try:
                result = future.result()
                if 'message' in result:
                    already_evaluated_messages.append(result['message'])
                results[submission_id] = result
            except Exception as e:
                logging.error(f"Error processing submission ID {submission_id}: {str(e)}")
                results[submission_id] = {"error": str(e)}

    return jsonify({'results': results, 'already_evaluated_messages': already_evaluated_messages}), 200

if __name__ == '__main__':
    app.run(debug=True)
